"""Document extractors for PDF, DOCX, TXT, and Markdown files.

Provides a unified extraction output structure containing titles, page divisions,
paragraphs, metadata, and statistics.
"""

from typing import Any

import docx
import pypdf
import re
from pydantic import BaseModel, Field

def sanitize_extracted_text(raw_text: str) -> str:
    # Strip out standalone page numbers (e.g., "Page 1", "Page 23")
    cleaned = re.sub(r'(?im)^Page\s+\d+\s*$', '', raw_text)
    # Remove rogue structural artifacts or excessive dots
    cleaned = re.sub(r'\.{4,}', '...', cleaned)
    # Replace 3 or more consecutive newlines with a clean double newline
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()


class ExtractionResult(BaseModel):
    """Unified structure returned by all document extractors."""

    title: str = Field(..., description="Document title or filename")
    pages: list[str] = Field(
        ...,
        description="Text content split by page (1-indexed internally as pages[0] = page 1)",
    )
    paragraphs: list[str] = Field(
        ..., description="Logical paragraphs extracted from document"
    )
    metadata: dict[str, Any] = Field(
        default_factory=dict, description="Raw file metadata"
    )
    statistics: dict[str, Any] = Field(
        default_factory=dict,
        description="Ingestion statistics (page_count, paragraph_count, char_count, word_count)",
    )


class BaseExtractor:
    """Base class for all document extractors."""

    @staticmethod
    def _compute_stats(pages: list[str], paragraphs: list[str]) -> dict[str, Any]:
        """Compute basic document statistics."""
        full_text = " ".join(pages)
        words = full_text.split()
        return {
            "page_count": len(pages),
            "paragraph_count": len(paragraphs),
            "char_count": len(full_text),
            "word_count": len(words),
        }


class PDFExtractor(BaseExtractor):
    """Extracts text page-by-page from PDF files, preferring PyMuPDF4LLM for rich Markdown."""

    @classmethod
    def extract(cls, file_content: bytes, filename: str) -> ExtractionResult:
        import logging
        logger = logging.getLogger("app")
        try:
            import fitz
            import pymupdf4llm
            
            doc = fitz.Document(stream=file_content, filetype="pdf")
            
            pages = []
            paragraphs = []
            metadata = {}
            
            if doc.metadata:
                metadata = {k: str(v) for k, v in doc.metadata.items() if v}
                
            title = metadata.get("title", filename)
            
            try:
                from langchain_text_splitters import RecursiveCharacterTextSplitter
            except ImportError:
                from langchain.text_splitter import RecursiveCharacterTextSplitter
                
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=800,
                chunk_overlap=150,
                separators=["\n\n", "\n", ".", " ", ""]
            )
            
            # Extract page by page to preserve pagination for chunks
            for page_num in range(len(doc)):
                try:
                    page_md = pymupdf4llm.to_markdown(doc=doc, pages=[page_num])
                    pages.append(page_md)
                    
                    if page_md.strip():
                        sanitized_text = sanitize_extracted_text(page_md)
                        page_paras = text_splitter.split_text(sanitized_text)
                        paragraphs.extend([p.strip() for p in page_paras if p.strip()])
                except Exception as e:
                    logger.warning(f"Failed to parse page {page_num} using pymupdf4llm: {e}")
                    pages.append("")
            
            stats = cls._compute_stats(pages, paragraphs)
            return ExtractionResult(
                title=title,
                pages=pages,
                paragraphs=paragraphs,
                metadata=metadata,
                statistics=stats,
            )
            
        except Exception as e:
            print(f"PyMuPDF4LLM IMPORT ERROR: {type(e).__name__} - {e}")
            logger.warning(f"PyMuPDF4LLM not installed or failed to import. Falling back to legacy PyPDF extractor. Error: {str(e)}")
            # Fallback to PyPDF
            import io
            import pypdf
            
            pdf_file = io.BytesIO(file_content)
            reader = pypdf.PdfReader(pdf_file)

            pages = []
            paragraphs = []
            metadata = {}

            if reader.metadata:
                for k, v in reader.metadata.items():
                    if v:
                        key = k.lstrip("/")
                        metadata[key] = str(v)

            title = metadata.get("Title", filename)

            try:
                from langchain_text_splitters import RecursiveCharacterTextSplitter
            except ImportError:
                from langchain.text_splitter import RecursiveCharacterTextSplitter
                
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=800,
                chunk_overlap=150,
                separators=["\n\n", "\n", ".", " ", ""]
            )

            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text = page.extract_text() or ""
                pages.append(text)

                if text.strip():
                    sanitized_text = sanitize_extracted_text(text)
                    page_paras = text_splitter.split_text(sanitized_text)
                    paragraphs.extend([p.strip() for p in page_paras if p.strip()])

            stats = cls._compute_stats(pages, paragraphs)
            return ExtractionResult(
                title=title,
                pages=pages,
                paragraphs=paragraphs,
                metadata=metadata,
                statistics=stats,
            )


class DOCXExtractor(BaseExtractor):
    """Extracts text from DOCX files using python-docx."""

    @classmethod
    def extract(cls, file_content: bytes, filename: str) -> ExtractionResult:
        import io
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)

        paragraphs = []
        pages = []
        metadata = {}

        # Extract metadata
        props = doc.core_properties
        if props.title:
            metadata["title"] = props.title
        if props.author:
            metadata["author"] = props.author

        title = metadata.get("title", filename)

        # Extract paragraph text
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())

        # Extract tables as markdown
        for table in doc.tables:
            if not table.rows:
                continue
                
            header_cells = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
            if not any(header_cells):
                continue
                
            paragraphs.append("| " + " | ".join(header_cells) + " |")
            paragraphs.append("|" + "|".join(["---"] * len(header_cells)) + "|")
            
            for row in table.rows[1:]:
                row_cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                paragraphs.append("| " + " | ".join(row_cells) + " |")

        # Since Word files lack hard pagination, group paragraphs into pages of ~3000 chars
        full_text = "\n\n".join(paragraphs)
        page_size = 3000
        if full_text:
            pages = [full_text[i:i + page_size] for i in range(0, len(full_text), page_size)]
        else:
            pages = [""]

        stats = cls._compute_stats(pages, paragraphs)
        return ExtractionResult(
            title=title,
            pages=pages,
            paragraphs=paragraphs,
            metadata=metadata,
            statistics=stats,
        )


class TXTExtractor(BaseExtractor):
    """Extracts text from plain text files."""

    @classmethod
    def extract(cls, file_content: bytes, filename: str) -> ExtractionResult:
        text = file_content.decode("utf-8", errors="ignore")
        pages = [text]
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        stats = cls._compute_stats(pages, paragraphs)
        return ExtractionResult(
            title=filename,
            pages=pages,
            paragraphs=paragraphs,
            metadata={"filename": filename},
            statistics=stats,
        )


class MarkdownExtractor(BaseExtractor):
    """Extracts text and title from Markdown files."""

    @classmethod
    def extract(cls, file_content: bytes, filename: str) -> ExtractionResult:
        text = file_content.decode("utf-8", errors="ignore")
        pages = [text]
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        # Try to find heading title (# Title)
        title = filename
        import re
        match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
        if match:
            title = match.group(1).strip()

        stats = cls._compute_stats(pages, paragraphs)
        return ExtractionResult(
            title=title,
            pages=pages,
            paragraphs=paragraphs,
            metadata={"filename": filename},
            statistics=stats,
        )
